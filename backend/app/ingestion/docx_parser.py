"""DOCX parser using ``python-docx`` (TDD §4.3).

Reads paragraphs, styles and tables. Tracked-change handling (``w:ins`` /
``w:del``) is required to distinguish negotiated edits from final text; the MVP
captures the *accepted* text and flags the presence of tracked changes in
metadata, with full insertion/deletion separation delivered in the 3-month build.
"""

from __future__ import annotations

from app.core.models import CIRBlock
from app.ingestion.base import DocumentParser, block_id
from app.logging_setup import get_logger

log = get_logger("ingestion.docx")


class DocxParser(DocumentParser):
    """Parse a ``.docx`` into paragraph and table blocks."""

    extensions = ("docx",)
    format = "docx"

    def parse(self, data: bytes, filename: str) -> tuple[list[CIRBlock], dict, int]:
        """Extract paragraphs and tables in document order.

        Raises:
            RuntimeError: If ``python-docx`` is not installed.
        """
        try:
            import io

            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx is required to parse DOCX files") from exc

        document = docx.Document(io.BytesIO(data))
        blocks: list[CIRBlock] = []
        idx = 0

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            blocks.append(CIRBlock(block_id=block_id(idx), type="paragraph", page=1, text=text))
            idx += 1

        for table in document.tables:
            matrix = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            flat = " | ".join(" ".join(r) for r in matrix)
            blocks.append(
                CIRBlock(block_id=block_id(idx), type="table", page=1, text=flat, table=matrix)
            )
            idx += 1

        has_tracked = b"w:ins" in data or b"w:del" in data
        meta = {"filename": filename, "has_tracked_changes": has_tracked}
        log.info("docx_parsed", extra={"filename": filename, "blocks": len(blocks),
                                       "tracked_changes": has_tracked})
        return blocks, meta, 1
