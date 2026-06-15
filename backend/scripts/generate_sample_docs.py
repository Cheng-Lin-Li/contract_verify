"""Generate PDF and DOCX versions of the sample playbook and standard-terms files.

Produces four files:
  samples/playbook/playbook.docx
  samples/playbook/playbook.pdf
  samples/standard_terms/standard_terms.docx
  samples/standard_terms/standard_terms.pdf

Run from the repo root:
    python backend/scripts/generate_sample_docs.py
"""

from __future__ import annotations

import sys
from pathlib import Path

# Resolve repo root and put backend/ on sys.path so app imports work.
REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT / "backend"))


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

PLAYBOOK_META = {
    "title": "Company Contract Playbook",
    "subtitle": "Layer 2 — Approved Clause Positions",
    "version": "v1.0",
    "description": (
        "This playbook defines the company's approved positions on key contract terms. "
        "Every incoming or outgoing agreement must be verified against these positions "
        "before signature. Items marked MUST HAVE are non-negotiable; items marked "
        "MUST NOT HAVE are absolute restrictions; PREFERRED items represent the "
        "company's standard ask but may be negotiated."
    ),
}

PLAYBOOK_SECTIONS = [
    {
        "heading": "1. Liability Cap",
        "body": (
            "Aggregate liability must be capped at the total fees paid under the agreement. "
            "No contract may expose the company to liability in excess of the fees actually "
            "received or paid under that specific agreement. This cap applies to all claims "
            "in aggregate, including breach of contract, tort, and statutory causes of action, "
            "whether or not the party has been advised of the possibility of such damages. "
            "The liability cap must appear in a clearly labelled Limitation of Liability "
            "section and must reference the specific contract value."
        ),
    },
    {
        "heading": "2. Uncapped Liability",
        "body": (
            "The agreement must not accept unlimited or uncapped liability for either party. "
            "Clauses that expose either party to unbounded financial liability -- including "
            "indemnities without a financial ceiling, consequential-damage waivers that "
            "operate in only one direction, or representations with no cap -- are prohibited. "
            "Any clause that effectively eliminates the liability cap for a broad category "
            "of claims must be redlined and flagged for attorney review before execution."
        ),
    },
    {
        "heading": "3. Confidentiality",
        "body": (
            "Confidentiality obligations should be mutual and survive termination. "
            "The preferred position is a mutual non-disclosure clause that binds both "
            "parties equally, with obligations surviving the expiry or termination of "
            "the agreement for a minimum period of three years. One-sided "
            "confidentiality clauses that bind only the company are acceptable only where "
            "the counterparty shares no confidential information. Any survival period of "
            "less than two years must be flagged for review. Confidential information "
            "should be defined broadly to include technical data, business plans, pricing, "
            "and customer information."
        ),
    },
    {
        "heading": "4. Payment Terms",
        "body": (
            "Payment terms shall not exceed net-45 days from invoice date. "
            "All agreements must specify a payment period no longer than forty-five "
            "calendar days from the date of a valid invoice. Agreements that specify "
            "net-60, net-90, or longer terms require Finance department pre-approval and "
            "must be documented in the deal record. Late-payment interest at the statutory "
            "rate or 1.5 percent per month, whichever is higher, should be included. "
            "Milestone-based payment schedules are permitted provided each milestone "
            "payment is individually due within net-45 of the relevant milestone "
            "acceptance notice."
        ),
    },
]

STANDARD_TERMS_META = {
    "title": "Standard Contract Terms Library",
    "subtitle": "Layer 3 — Market-Standard Baseline for Services Agreements",
    "version": "v1.0",
    "description": (
        "This library defines the market-standard terms expected to be present in every "
        "services agreement. These terms represent baseline protections and obligations "
        "that a well-formed commercial contract should include, regardless of the specific "
        "deal requirements or company playbook positions. Missing terms are flagged "
        "automatically during verification and routed for attorney review."
    ),
}

STANDARD_TERMS_SECTIONS = [
    {
        "heading": "1. Governing Law and Jurisdiction",
        "body": (
            "Every services agreement must specify the governing law and jurisdiction. "
            "The clause must identify the state or jurisdiction whose laws govern the "
            "interpretation and enforcement of the agreement, the courts of competent "
            "jurisdiction for dispute resolution, and whether the parties consent to "
            "personal jurisdiction in those courts. "
            "The clause must identify a single, definitive legal system and must not "
            "leave governing law to be implied by the domicile of either party. "
            "Choice-of-law provisions that select a jurisdiction without specifying "
            "the dispute resolution forum are incomplete."
        ),
    },
    {
        "heading": "2. Limitation of Liability",
        "body": (
            "A bilateral limitation of liability clause must appear in every services agreement, "
            "capping aggregate damages at the fees paid or payable under the agreement during "
            "the preceding twelve months. The clause should also exclude consequential, "
            "incidental, special, and punitive damages for both parties. Standard carve-outs "
            "include fraud, gross negligence, wilful misconduct, and death or personal injury "
            "caused by negligence. IP indemnities with a financial ceiling are also "
            "standard and should be captured separately."
        ),
    },
    {
        "heading": "3. Mutual Confidentiality",
        "body": (
            "A mutual confidentiality clause protecting each party's confidential information "
            "is standard in all services agreements. The clause should define confidential "
            "information broadly, specify permitted disclosures to employees and advisors, "
            "impose a duty of care no less than the receiving party uses for its own "
            "confidential information, and survive termination for at least two years. "
            "Carve-outs for publicly available information and information received from "
            "third parties without restriction are standard."
        ),
    },
    {
        "heading": "4. Indemnification",
        "body": (
            "A mutual indemnification clause allocating liability for third-party claims "
            "is a standard term in services agreements. The clause must address "
            "indemnification for IP infringement claims arising from each party's own "
            "materials, for claims arising from gross negligence or wilful misconduct, "
            "and must include a process for tendering defence and requiring cooperation. "
            "Unlimited indemnities without a financial cap are non-standard and must be "
            "flagged. Asymmetric IP indemnities where the vendor indemnifies the customer "
            "are common in SaaS and services agreements."
        ),
    },
    {
        "heading": "5. Data Protection and Deletion",
        "body": (
            "Any services agreement under which personal data or proprietary data is "
            "processed must address data protection and data deletion obligations. "
            "The agreement must specify the nature and purpose of data processing, "
            "the technical and organisational security measures in place, the obligation "
            "to notify of a data breach within 72 hours for personal data under "
            "GDPR-aligned frameworks, and certified deletion or return of all data "
            "within thirty days of termination or expiry. Agreements involving personal "
            "data of EU or UK data subjects must reference the applicable transfer "
            "mechanism such as Standard Contractual Clauses. "
            "Absence of a data-deletion clause is a mandatory escalation item."
        ),
    },
]


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _p(text: str) -> str:
    """Normalize text to latin-1-safe characters for fpdf core fonts."""
    return (
        text
        .replace("—", "--")   # em dash
        .replace("–", "-")    # en dash
        .replace("‘", "'").replace("’", "'")   # curly single quotes
        .replace("“", '"').replace("”", '"')   # curly double quotes
        .replace("•", "*")    # bullet
        .replace("●", "o")    # black circle (was ●)
        .replace("✕", "x")    # multiplication x (was ✕)
        .replace("◎", "o")    # bullseye (was ◎)
    )


def generate_docx(meta: dict, sections: list[dict], out_path: Path) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()

    # Title
    title = doc.add_heading(meta["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(meta["subtitle"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    ver = doc.add_paragraph(meta["version"])
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Description
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(meta["description"])
    doc.add_paragraph()

    # Sections
    for sec in sections:
        doc.add_heading(sec["heading"], level=1)
        doc.add_paragraph(sec["body"])
        doc.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  Wrote {out_path}")


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def generate_pdf(meta: dict, sections: list[dict], out_path: Path) -> None:
    from fpdf import FPDF

    class Doc(FPDF):
        def header(self):
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 6, _p(meta["title"]), align="R")
            self.ln(4)
            self.set_draw_color(200, 200, 200)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(3)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 6, f"Page {self.page_no()}", align="C")

    pdf = Doc()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin  # usable content width

    # Title block
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(20, 20, 80)
    pdf.multi_cell(W, 10, _p(meta["title"]), align="C")
    pdf.set_font("Helvetica", "I", 12)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(W, 7, _p(meta["subtitle"]), align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(130, 130, 130)
    pdf.multi_cell(W, 6, _p(meta["version"]), align="C")
    pdf.ln(6)

    # Divider
    pdf.set_draw_color(60, 60, 140)
    pdf.set_line_width(0.8)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(6)

    # Overview
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(20, 20, 80)
    pdf.cell(W, 8, "Overview", ln=True)
    pdf.set_line_width(0.3)
    pdf.set_draw_color(180, 180, 200)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(3)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(40, 40, 40)
    pdf.multi_cell(W, 5.5, _p(meta["description"]))
    pdf.ln(6)

    for sec in sections:
        # Force page break if less than 45mm remain
        if pdf.get_y() > pdf.h - pdf.b_margin - 45:
            pdf.add_page()

        # Section heading
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(20, 20, 80)
        pdf.cell(W, 8, _p(sec["heading"]), ln=True)
        pdf.set_line_width(0.3)
        pdf.set_draw_color(180, 180, 200)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

        # Body text
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(W, 5.5, _p(sec["body"]))
        pdf.ln(6)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))
    print(f"  Wrote {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    root = REPO_ROOT

    print("Generating playbook documents…")
    generate_docx(PLAYBOOK_META, PLAYBOOK_SECTIONS,
                  root / "samples" / "playbook" / "playbook.docx")
    generate_pdf(PLAYBOOK_META, PLAYBOOK_SECTIONS,
                 root / "samples" / "playbook" / "playbook.pdf")

    print("Generating standard-terms documents…")
    generate_docx(STANDARD_TERMS_META, STANDARD_TERMS_SECTIONS,
                  root / "samples" / "standard_terms" / "standard_terms.docx")
    generate_pdf(STANDARD_TERMS_META, STANDARD_TERMS_SECTIONS,
                 root / "samples" / "standard_terms" / "standard_terms.pdf")

    print("Done.")


if __name__ == "__main__":
    main()
